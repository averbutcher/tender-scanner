import asyncio
import json
import os
import secrets
import sys
import tempfile
import time
from pathlib import Path
from typing import Optional

import anthropic as _anthropic
import bcrypt
import yaml
from fastapi import Cookie, Depends, FastAPI, File, Form, HTTPException, Query, Request, Response, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from yaml.loader import SafeLoader
from dotenv import load_dotenv

sys.path.insert(0, str(Path(__file__).parent))

from analyzer import analyze_tender, distill_knowledge, SYSTEM_PROMPT
from engine import load_config, save_config, parse_message, parse_excel, compare, export_to_excel, find_suspicious_lines
from scraper import Tender, _extract_id_from_url, _extract_pdf_text_from_bytes, fetch_tender_detail, fetch_tender_list
from state import filter_new, load_seen, save_seen

BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / ".env")

_SIGNER = URLSafeTimedSerializer(os.environ.get("SESSION_SECRET", "et-tools-2024-change-me"))
SETTINGS_FILE = BASE_DIR / "data" / "settings.json"
KNOWLEDGE_FILE = BASE_DIR / "data" / "shared" / "knowledge.json"
_excel_cache: dict[str, bytes] = {}

app = FastAPI(title="Electra Target Tools")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

@app.on_event("startup")
async def _init_dirs():
    for d in ["data/shared", "data/users"]:
        (BASE_DIR / d).mkdir(parents=True, exist_ok=True)


# ── Session helpers ───────────────────────────────────────────────────────────

def _sign(username: str) -> str:
    return _SIGNER.dumps(username)

def _unsign(token: Optional[str]) -> Optional[str]:
    if not token:
        return None
    try:
        return _SIGNER.loads(token, max_age=30 * 86400)
    except (BadSignature, SignatureExpired, Exception):
        return None

def auth(et_session: Optional[str] = Cookie(default=None)) -> str:
    u = _unsign(et_session)
    if not u:
        raise HTTPException(401, "לא מחובר")
    return u

def _load_users() -> dict:
    with open(BASE_DIR / "users.yaml", encoding="utf-8") as f:
        return yaml.load(f, SafeLoader)


# ── Data helpers ──────────────────────────────────────────────────────────────

def _udir(u: str) -> Path:
    p = BASE_DIR / "data" / "users" / u
    p.mkdir(parents=True, exist_ok=True)
    return p

def _rj(path: Path, default):
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            pass
    return default

def _wj(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def _load_settings() -> dict:
    d = {
        "scraper": {
            "base_url": "https://mr.gov.il/ilgstorefront/he/search/?q=:relevance&inContract=false",
            "page_load_timeout_ms": 30000, "max_tenders_per_run": 50, "days_back": 7,
        },
        "budget": {"min_annual_ils": 500000, "max_annual_ils": 50000000},
        "industries": ["ניקיון", "אחזקה", "שמירה ואבטחה", "כוח אדם", "שירותי עזר", "קייטרינג"],
        "labor_costs": {"simple_monthly_ils": 6500, "simple_hourly_ils": 35, "social_expense_multiplier": 1.25},
    }
    if not SETTINGS_FILE.exists():
        _wj(SETTINGS_FILE, d)
        return d
    loaded = _rj(SETTINGS_FILE, d)
    for k, v in d.items():
        loaded.setdefault(k, v)
    return loaded

def _load_knowledge() -> list:
    return _rj(KNOWLEDGE_FILE, [])

def _save_knowledge(g: list):
    KNOWLEDGE_FILE.parent.mkdir(parents=True, exist_ok=True)
    _wj(KNOWLEDGE_FILE, g)

def _append_history(result: dict, u: str):
    h = _rj(_udir(u) / "history.json", [])
    if result["tender_id"] not in {r["tender_id"] for r in h}:
        h.append(result)
        _wj(_udir(u) / "history.json", h)

def _client() -> _anthropic.Anthropic:
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        raise HTTPException(500, "ANTHROPIC_API_KEY לא מוגדר")
    return _anthropic.Anthropic(api_key=key)


# ── Auth routes ───────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def index():
    return HTMLResponse((BASE_DIR / "static" / "index.html").read_text(encoding="utf-8"))

@app.post("/api/login")
async def login(resp: Response, username: str = Form(...), password: str = Form(...)):
    data = _load_users()
    user = data["credentials"]["usernames"].get(username)
    if not user or not bcrypt.checkpw(password.encode(), user["password"].encode()):
        raise HTTPException(401, "שם משתמש או סיסמה שגויים")
    resp.set_cookie("et_session", _sign(username), max_age=30 * 86400, httponly=True, samesite="lax")
    return {"username": username, "name": user.get("name", username),
            "role": user.get("role", "user"), "apps": user.get("apps", ["tender_scanner", "shift_comparison"])}

@app.post("/api/logout")
async def logout(resp: Response):
    resp.delete_cookie("et_session")
    return {"ok": True}

@app.get("/api/me")
async def me(u: str = Depends(auth)):
    users = _load_users()["credentials"]["usernames"]
    info = users.get(u, {})
    return {"username": u, "name": info.get("name", u),
            "role": info.get("role", "user"), "apps": info.get("apps", ["tender_scanner", "shift_comparison"])}


# ── Settings & knowledge ──────────────────────────────────────────────────────

@app.get("/api/settings")
async def get_settings(_: str = Depends(auth)):
    return _load_settings()

@app.post("/api/settings")
async def post_settings(body: dict, _: str = Depends(auth)):
    _wj(SETTINGS_FILE, body)
    return {"ok": True}

@app.get("/api/knowledge")
async def get_knowledge(_: str = Depends(auth)):
    return _load_knowledge()

@app.post("/api/knowledge")
async def post_knowledge(body: dict, _: str = Depends(auth)):
    _save_knowledge(body.get("guidelines", []))
    return {"ok": True}

@app.delete("/api/knowledge/{idx}")
async def del_knowledge(idx: int, _: str = Depends(auth)):
    k = _load_knowledge()
    if 0 <= idx < len(k):
        k.pop(idx)
        _save_knowledge(k)
    return k


# ── Tender data ───────────────────────────────────────────────────────────────

@app.get("/api/history")
async def get_history(u: str = Depends(auth)):
    return _rj(_udir(u) / "history.json", [])

@app.patch("/api/history/{tid}")
async def patch_history(tid: str, body: dict, u: str = Depends(auth)):
    path = _udir(u) / "history.json"
    h = _rj(path, [])
    for entry in h:
        if entry.get("tender_id") == tid:
            entry.update(body)
            break
    _wj(path, h)
    return {"ok": True}

@app.delete("/api/history/{tid}")
async def delete_history(tid: str, u: str = Depends(auth)):
    path = _udir(u) / "history.json"
    h = _rj(path, [])
    h = [r for r in h if r.get("tender_id") != tid]
    _wj(path, h)
    # also remove from seen so it can be re-scanned
    seen_path = _udir(u) / "seen.json"
    seen = load_seen(seen_path)
    seen.discard(tid)
    save_seen(seen, seen_path)
    return {"ok": True}

@app.get("/api/last-scan")
async def get_last_scan(u: str = Depends(auth)):
    return _rj(_udir(u) / "last_scan.json", [])

@app.get("/api/favorites")
async def get_favorites(u: str = Depends(auth)):
    return _rj(_udir(u) / "favorites.json", [])

@app.post("/api/favorites/{tid}")
async def toggle_fav(tid: str, u: str = Depends(auth)):
    favs: list = _rj(_udir(u) / "favorites.json", [])
    if tid in favs:
        favs.remove(tid)
    else:
        favs.append(tid)
    _wj(_udir(u) / "favorites.json", favs)
    return favs


# ── Scan (SSE) ────────────────────────────────────────────────────────────────

@app.get("/api/scan-test")
async def scan_test(u: str = Depends(auth)):
    import traceback
    try:
        settings = _load_settings()
        from playwright.async_api import async_playwright
        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True)
            context = await browser.new_context(locale="he-IL", extra_http_headers={"Accept-Language": "he-IL,he;q=0.9"})
            page = await context.new_page()
            await page.goto(settings["scraper"]["base_url"], timeout=30000, wait_until="domcontentloaded")
            await page.wait_for_timeout(3000)
            final_url = page.url
            items = await page.query_selector_all("div.result-container")
            # check show-more button with various selectors
            show_more = None
            show_more_sel = None
            for sel in ["button.show-more-button","button:has-text('הצג עוד')","a:has-text('הצג עוד')",".show-more","[class*='show-more']"]:
                el = await page.query_selector(sel)
                if el:
                    show_more = await el.inner_text()
                    show_more_sel = sel
                    break
            # get first 3 titles
            titles = []
            for item in items[:3]:
                el = await item.query_selector("h2")
                if el: titles.append(await el.inner_text())
            await browser.close()
            return {"final_url": final_url, "items_found": len(items), "show_more_btn": show_more, "show_more_sel": show_more_sel, "sample_titles": titles}
    except Exception as e:
        return {"ok": False, "error": str(e), "trace": traceback.format_exc()[-1000:]}

@app.get("/api/scan")
async def scan(u: str = Depends(auth), skip_seen: bool = Query(True)):
    settings = _load_settings()

    async def gen():
        import traceback
        try:
            client = _client()
            loop = asyncio.get_running_loop()
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','msg':f'שגיאת אתחול: {traceback.format_exc()[-400:]}'})}\n\n"
            return

        yield f"data: {json.dumps({'type':'status','msg':'מתחבר ל-mr.gov.il...'})}\n\n"
        try:
            import traceback
            tender_list = await asyncio.wait_for(fetch_tender_list(settings), timeout=120)
        except asyncio.TimeoutError:
            yield f"data: {json.dumps({'type':'error','msg':'timeout — הסריקה לקחה יותר מ-120 שניות'})}\n\n"
            return
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','msg': traceback.format_exc()[-500:]})}\n\n"
            return

        seen_path = _udir(u) / "seen.json"
        seen = load_seen(seen_path) if skip_seen else set()
        new = filter_new(tender_list, seen) if skip_seen else tender_list

        yield f"data: {json.dumps({'type':'status','msg':f'נמצאו {len(tender_list)} מכרזים, {len(new)} חדשים לניתוח'})}\n\n"
        yield f"data: {json.dumps({'type':'count','total':len(tender_list),'new':len(new)})}\n\n"

        if not new:
            yield f"data: {json.dumps({'type':'complete','count':0,'total':len(tender_list)})}\n\n"
            return

        results = []
        for i, meta in enumerate(new):
            yield f"data: {json.dumps({'type':'progress','i':i+1,'total':len(new),'title':meta['title'][:80]})}\n\n"
            yield ": keepalive\n\n"

            try:
                tender = await asyncio.wait_for(fetch_tender_detail(meta, settings), timeout=90)
                tender.raw_metadata.update({"publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date","")})

                if not tender.pdf_text:
                    result = {
                        "tender_id": tender.tender_id, "title": tender.title, "url": tender.url,
                        "publisher": tender.publisher, "deadline": tender.deadline,
                        "publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date",""),
                        "has_pdf": False, "analysis": "NO_PDF",
                    }
                else:
                    knowledge = _load_knowledge()
                    result = await loop.run_in_executor(
                        None, lambda t=tender: analyze_tender(t, settings, client, knowledge=knowledge)
                    )
            except asyncio.TimeoutError:
                result = {
                    "tender_id": meta["tender_id"], "title": meta["title"], "url": meta["url"],
                    "publisher": "", "deadline": "",
                    "publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date",""),
                    "has_pdf": False, "analysis": "שגיאה: timeout",
                }
            except Exception as e:
                result = {
                    "tender_id": meta["tender_id"], "title": meta["title"], "url": meta["url"],
                    "publisher": "", "deadline": "",
                    "publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date",""),
                    "has_pdf": False, "analysis": f"שגיאה: {e}",
                }

            results.append(result)
            _append_history(result, u)
            seen.add(meta["tender_id"])
            if skip_seen:
                save_seen(seen, seen_path)
            _wj(_udir(u) / "last_scan.json", results)

            yield f"data: {json.dumps({'type':'result','data':result,'pct':(i+1)/len(new)})}\n\n"

        yield f"data: {json.dumps({'type':'complete','count':len(results),'total':len(tender_list)})}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream",
                              headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Analyze single URL ────────────────────────────────────────────────────────

@app.post("/api/analyze")
async def analyze_url(body: dict, u: str = Depends(auth)):
    url = body.get("url", "").strip()
    if not url:
        raise HTTPException(400, "URL נדרש")
    settings = _load_settings()
    tid = _extract_id_from_url(url)
    meta = {"tender_id": tid, "title": url, "url": url, "publish_date": "", "update_date": ""}
    try:
        tender = await asyncio.wait_for(fetch_tender_detail(meta, settings), timeout=90)
        if not tender.title or tender.title.startswith("http"):
            tender.title = f"מכרז {tid}"
        tender.raw_metadata.update({"publish_date": "", "update_date": ""})
        if not tender.pdf_text:
            result = {
                "tender_id": tid, "title": tender.title, "url": url,
                "publisher": tender.publisher, "deadline": tender.deadline,
                "publish_date": "", "update_date": "", "has_pdf": False, "analysis": "NO_PDF",
            }
        else:
            client = _client()
            loop = asyncio.get_running_loop()
            knowledge = _load_knowledge()
            result = await loop.run_in_executor(
                None, lambda t=tender: analyze_tender(t, settings, client, knowledge=knowledge)
            )
        _append_history(result, u)
        return result
    except asyncio.TimeoutError:
        raise HTTPException(408, "timeout")
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Analyze uploaded PDF ──────────────────────────────────────────────────────

@app.post("/api/analyze-pdf")
async def analyze_pdf_upload(pdf: UploadFile = File(...), u: str = Depends(auth)):
    body = await pdf.read()
    if not body:
        raise HTTPException(400, "קובץ ריק")
    pdf_text = _extract_pdf_text_from_bytes(body)
    if not pdf_text.strip():
        raise HTTPException(422, "לא ניתן לחלץ טקסט מהקובץ")
    settings = _load_settings()
    client = _client()
    knowledge = _load_knowledge()
    filename = pdf.filename or "מכרז"
    title = filename.replace(".pdf", "").replace("_", " ")
    tender = Tender(
        tender_id=f"pdf_{int(time.time())}",
        title=title,
        url="",
        publisher="",
        deadline="",
        raw_metadata={"publish_date": "", "update_date": ""},
        pdf_text=pdf_text,
    )
    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(
        None, lambda: analyze_tender(tender, settings, client, knowledge=knowledge)
    )
    _append_history(result, u)
    return result


# ── Chat ──────────────────────────────────────────────────────────────────────

@app.post("/api/chat")
async def chat(body: dict, _: str = Depends(auth)):
    tender = body.get("tender", {})
    history = body.get("history", [])
    client = _client()
    system = f"""אתה יועץ עסקי שניתח מכרז עבור Electra Target.
כותרת: {tender.get('title','')}
מפרסם: {tender.get('publisher','לא ידוע')}
מועד הגשה: {tender.get('deadline','לא צוין')}
ניתוח: {tender.get('analysis','')}
ענה בעברית בלבד. היה ממוקד ותמציתי."""
    msgs = [{"role": m["role"], "content": m["content"]} for m in history]
    loop = asyncio.get_running_loop()
    try:
        resp = await loop.run_in_executor(
            None,
            lambda: client.messages.create(
                model="claude-sonnet-4-6", max_tokens=1024, system=system, messages=msgs
            )
        )
        return {"reply": resp.content[0].text}
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Re-analyze with clarification answers ─────────────────────────────────────

@app.post("/api/tender/reanalyze-with-answers")
async def reanalyze_with_answers(
    answers: UploadFile = File(...),
    tender_data: str = Form(...),
    u: str = Depends(auth)
):
    import json as _json
    td = _json.loads(tender_data)
    answers_text = (await answers.read()).decode("utf-8", errors="ignore")
    settings = _load_settings()
    client = _client()
    knowledge = _load_knowledge()
    tender = Tender(
        tender_id=td.get("tender_id",""),
        title=td.get("title",""),
        url=td.get("url",""),
        publisher=td.get("publisher",""),
        deadline=td.get("deadline",""),
        raw_metadata={"publish_date": td.get("publish_date",""), "update_date": td.get("update_date","")},
        pdf_text=td.get("pdf_text",""),
    )
    feedback = [f"תשובות הבהרה שהתקבלו:\n{answers_text}"]
    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(
        None, lambda: analyze_tender(tender, settings, client, knowledge=knowledge, session_feedback=feedback)
    )
    _append_history(result, u)
    return result


# ── Generate clarification questions ──────────────────────────────────────────

@app.post("/api/tender/generate-questions")
async def generate_questions(body: dict, _: str = Depends(auth)):
    r = body
    client = _client()
    knowledge = _load_knowledge()
    prompt = f"""אתה יועץ עסקי של Electra Target. קראת את הניתוח הבא של מכרז ממשלתי.

כותרת: {r.get('title','')}
מפרסם: {r.get('publisher','')}
מועד הגשה: {r.get('deadline','')}
ניתוח: {r.get('analysis','')}

צור את כל שאלות ההבהרה שיש לשלוח למפרסם המכרז — כמה שצריך, עד 50 שאלות לכל היותר.
כלול כל שאלה חשובה שעולה מהמסמך. אל תגביל את עצמך למספר קבוע.

פלט כל שאלה בפורמט הבא (עמודה מופרדת ב-|):
מספר|עמוד|סעיף|שאלה

- עמוד: מספר עמוד רלוונטי במסמך אם ידוע, אחרת: כללי
- סעיף: מספר סעיף רלוונטי אם ידוע, אחרת: כללי
- שאלה: טקסט השאלה בעברית

דוגמה:
1|כללי|כללי|האם נדרש רישיון עסק?
2|5|3.2|מה תקופת האחריות על הציוד?

כתוב את כל השאלות בפורמט זה בלבד, ללא כותרות נוספות."""

    system = SYSTEM_PROMPT
    if knowledge:
        system += "\n\nתובנות שנצברו:\n" + "\n".join(f"- {g}" for g in knowledge)

    loop = asyncio.get_running_loop()
    try:
        resp = await loop.run_in_executor(
            None,
            lambda: client.messages.create(
                model="claude-sonnet-4-6", max_tokens=4096, system=system,
                messages=[{"role": "user", "content": prompt}]
            )
        )
        return {"questions": resp.content[0].text}
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Export Excel (financial analysis) ─────────────────────────────────────────

@app.post("/api/tender/export-excel")
async def export_excel(body: dict, _: str = Depends(auth)):
    import io, re
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    r = body
    analysis = r.get("analysis", "")

    # Split analysis into sections
    fin_keys = ["הערכת היקף כספי", "ערך שנתי", "בסיס לחישוב", "כוח אדם נדרש", "אורך חוזה", "המלצה", "אתגרים", "סיכונים"]
    sections: dict[str, list[str]] = {}
    current_key = None
    for line in analysis.splitlines():
        clean = line.strip().replace("**","").replace("#","").strip()
        if not clean: continue
        matched = next((k for k in fin_keys if k in clean), None)
        if matched:
            current_key = matched
            sections.setdefault(current_key, [])
        if current_key:
            sections[current_key].append(clean)

    wb = Workbook()
    ws = wb.active
    ws.title = "ניתוח פיננסי"
    ws.sheet_view.rightToLeft = True

    hfill = PatternFill("solid", fgColor="1E3A5F")
    sfill = PatternFill("solid", fgColor="2563EB")
    thfill= PatternFill("solid", fgColor="374151")
    afill = PatternFill("solid", fgColor="EEF3FA")
    wfill = PatternFill("solid", fgColor="FFFFFF")
    thin  = Border(left=Side(style='thin',color='CCCCCC'), right=Side(style='thin',color='CCCCCC'),
                   top=Side(style='thin',color='CCCCCC'),  bottom=Side(style='thin',color='CCCCCC'))

    def is_md_table(line):
        return line.startswith('|') and line.endswith('|')

    def is_separator(line):
        return re.fullmatch(r'[\|\-\s:]+', line) is not None

    def parse_md_row(line):
        return [c.strip() for c in line.strip('|').split('|')]

    def rtl_align(horizontal="right", center=False):
        return Alignment(horizontal="center" if center else horizontal,
                         vertical="center", wrap_text=True)

    def style_cell(c, bold=False, fill=None, center=False):
        c.font = Font(bold=bold, name="Arial", size=10)
        if fill: c.fill = fill
        c.border = thin
        c.alignment = rtl_align(center=center)

    def hrow(label, row, fill=None, ncols=2):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        c = ws.cell(row, 1, label)
        c.font = Font(bold=True, color="FFFFFF", size=12, name="Arial")
        c.fill = fill or hfill
        c.alignment = rtl_align()
        ws.row_dimensions[row].height = 24

    def drow(label, value, row, alt=False, ncols=2):
        fill = afill if alt else wfill
        c1 = ws.cell(row, 1, label or "")
        style_cell(c1, bold=bool(label), fill=fill)
        c2 = ws.cell(row, 2, str(value) if value else "")
        style_cell(c2, fill=fill)
        ws.row_dimensions[row].height = max(18, min(90, len(str(value or ""))//3+15))

    def write_section_lines(lines, start_row, max_cols):
        r = start_row
        alt = False
        i = 0
        while i < len(lines):
            line = lines[i]
            if is_md_table(line):
                # collect table block
                tbl_lines = []
                while i < len(lines) and (is_md_table(lines[i]) or is_separator(lines[i])):
                    tbl_lines.append(lines[i]); i += 1
                data_rows = [parse_md_row(l) for l in tbl_lines if not is_separator(l)]
                if not data_rows: continue
                ncols = max(len(row) for row in data_rows)
                # set/extend column widths
                for ci in range(1, ncols+1):
                    col_letter = ws.cell(r, ci).column_letter
                    ws.column_dimensions[col_letter].width = max(
                        ws.column_dimensions[col_letter].width, 18)
                for ri, dr in enumerate(data_rows):
                    is_hdr = ri == 0
                    row_fill = thfill if is_hdr else (afill if ri%2==0 else wfill)
                    for ci, val in enumerate(dr):
                        c = ws.cell(r, ci+1, val)
                        style_cell(c, bold=is_hdr, fill=row_fill,
                                   center=(ci < len(dr)-1))
                        if is_hdr: c.font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
                    ws.row_dimensions[r].height = 20
                    r += 1
            else:
                clean = line.replace("**","").replace("#","").strip()
                if clean and not is_separator(clean):
                    c = ws.cell(r, 1, clean)
                    style_cell(c, fill=(afill if alt else wfill))
                    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=max_cols)
                    ws.row_dimensions[r].height = max(18, min(60, len(clean)//5+15))
                    alt = not alt
                    r += 1
                i += 1
        return r

    # Auto-detect column count from tables in analysis
    max_cols = 2
    for line in analysis.splitlines():
        if is_md_table(line.strip()):
            max_cols = max(max_cols, line.count('|') - 1)
    max_cols = min(max_cols, 8)
    for ci in range(1, max_cols+1):
        letter = ws.cell(1, ci).column_letter
        ws.column_dimensions[letter].width = 20
    ws.column_dimensions[ws.cell(1,1).column_letter].width = 30

    row = 1
    hrow(f"ניתוח פיננסי: {r.get('title','')}", row, ncols=max_cols); row += 1
    for label, val, alt in [
        ("מפרסם", r.get("publisher",""), False),
        ("מועד הגשה", r.get("deadline",""), True),
        ("תאריך פרסום", r.get("publish_date",""), False),
    ]:
        drow(label, val, row, alt, ncols=max_cols); row += 1
    row += 1

    finance_order = ["הערכת היקף כספי", "ערך שנתי", "בסיס לחישוב", "כוח אדם נדרש", "אורך חוזה", "המלצה", "אתגרים", "סיכונים"]
    hrow("ניתוח פיננסי מפורט", row, sfill, ncols=max_cols); row += 1
    shown = set()
    for key in finance_order:
        lines = sections.get(key)
        if not lines or key in shown: continue
        shown.add(key)
        # Section sub-header
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=max_cols)
        c = ws.cell(row, 1, key)
        c.font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
        c.fill = PatternFill("solid", fgColor="2B5C9E")
        c.alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 20; row += 1
        row = write_section_lines(lines[1:], row, max_cols)  # skip key header line

    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    tid = r.get("tender_id", "tender")
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''%D7%A4%D7%99%D7%A0%D7%A0%D7%A1%D7%99_{tid}.xlsx"})


# ── Export Word (full analysis) ────────────────────────────────────────────────

@app.post("/api/tender/export-word")
async def export_word(body: dict, _: str = Depends(auth)):
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r = body
    analysis  = r.get("analysis", "")
    questions = r.get("questions", "")
    knowledge = _load_knowledge()

    doc = Document()
    # Set document-level RTL
    sectPr = doc.sections[0]._sectPr
    sectPr.append(OxmlElement('w:bidi'))

    def set_rtl_para(par):
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = par._p.get_or_add_pPr()
        b = OxmlElement('w:bidi'); b.set(qn('w:val'), '1'); pPr.append(b)
        jc = OxmlElement('w:jc');  jc.set(qn('w:val'), 'right'); pPr.append(jc)

    def set_rtl_run(run):
        run.font.cs_name = "Arial"
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)
        rPr.append(OxmlElement('w:cs'))

    def h(text, level=1):
        par = doc.add_heading(text, level=level)
        set_rtl_para(par)
        for run in par.runs:
            run.font.name = "Arial"
            set_rtl_run(run)
            if level == 1:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        return par

    from docx.shared import Cm
    from docx.oxml import OxmlElement as OE

    def add_run_inline(par, text):
        """Add run with inline **bold** parsing."""
        import re
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            if not part: continue
            run = par.add_run(part)
            run.font.name = "Arial"; run.font.size = Pt(11)
            run.bold = (i % 2 == 1)
            set_rtl_run(run)

    def p(text, bold=False, bullet=False, size=11):
        par = doc.add_paragraph()
        set_rtl_para(par)
        if bullet:
            pPr = par._p.get_or_add_pPr()
            ind = OE('w:ind'); ind.set(qn('w:right'), '360'); pPr.append(ind)
        if bold:
            run = par.add_run(str(text))
            run.font.name = "Arial"; run.font.cs_name = "Arial"
            run.font.size = Pt(size); run.bold = True
            set_rtl_run(run)
        else:
            add_run_inline(par, str(text))
        return par

    def render_md(text):
        """Render markdown text into the Word document."""
        import re
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            # Heading
            hm = re.match(r'^(#{1,3})\s+(.*)', stripped)
            if hm:
                level = min(len(hm.group(1)) + 1, 3)
                h(hm.group(2).replace('**',''), level)
                i += 1; continue
            # Table
            if stripped.startswith('|') and stripped.endswith('|'):
                tbl_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    tbl_lines.append(lines[i].strip()); i += 1
                data_rows = [r for r in tbl_lines if not re.fullmatch(r'[\|\-\s:]+', r)]
                if not data_rows: continue
                parsed = [[c.strip() for c in r.strip('|').split('|')] for r in data_rows]
                ncols = max(len(r) for r in parsed)
                tbl = doc.add_table(rows=len(parsed), cols=ncols)
                tbl.style = 'Table Grid'
                for ri, row in enumerate(parsed):
                    for ci, val in enumerate(row):
                        cell = tbl.rows[ri].cells[ci]
                        cell.paragraphs[0].clear()
                        cp = cell.paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        pPr = cp._p.get_or_add_pPr()
                        b = OE('w:bidi'); b.set(qn('w:val'),'1'); pPr.append(b)
                        run = cp.add_run(val)
                        run.font.name = "Arial"; run.font.size = Pt(10)
                        run.bold = (ri == 0)
                        if ri == 0: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        set_rtl_run(run)
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OE('w:shd')
                        shd.set(qn('w:fill'), '1E3A5F' if ri==0 else ('EEF3FA' if ri%2==0 else 'FFFFFF'))
                        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear')
                        tcPr.append(shd)
                continue
            # Bullet
            bm = re.match(r'^[-*+]\s+(.*)', stripped)
            if bm:
                p(f"• {bm.group(1)}", bullet=True); i += 1; continue
            # Numbered list
            nm = re.match(r'^\d+\.\s+(.*)', stripped)
            if nm:
                p(f"• {nm.group(1)}", bullet=True); i += 1; continue
            # Horizontal rule or separator — skip
            if re.fullmatch(r'[-_*]{2,}', stripped):
                i += 1; continue
            # Empty
            if not stripped:
                i += 1; continue
            # Normal
            p(stripped)
            i += 1

    h(f"ניתוח מכרז: {r.get('title','')}", 1)
    if r.get('publisher'):    p(f"מפרסם: {r['publisher']}")
    if r.get('publish_date'): p(f"תאריך פרסום: {r['publish_date']}")
    if r.get('deadline'):     p(f"מועד הגשה: {r['deadline']}")
    if r.get('update_date'):  p(f"תאריך עדכון: {r['update_date']}")
    doc.add_paragraph()

    render_md(analysis.replace('<','').replace('>',''))

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    tid = r.get("tender_id", "tender")
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''%D7%A0%D7%99%D7%AA%D7%95%D7%97_{tid}.docx"})


# ── Export questions to Word (table) ──────────────────────────────────────────

@app.post("/api/tender/export-questions-word")
async def export_questions_word(body: dict, _: str = Depends(auth)):
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r = body
    questions_raw = r.get("questions", "")

    doc = Document()
    sectPr = doc.sections[0]._sectPr
    sectPr.append(OxmlElement('w:bidi'))

    def set_rtl(par):
        pPr = par._p.get_or_add_pPr()
        b = OxmlElement('w:bidi'); b.set(qn('w:val'), '1'); pPr.append(b)
        jc = OxmlElement('w:jc');  jc.set(qn('w:val'), 'right'); pPr.append(jc)

    def set_rtl_run(run):
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)
        rPr.append(OxmlElement('w:cs'))
        run.font.cs_name = "Arial"

    title_p = doc.add_heading(f"שאלות הבהרה: {r.get('title','')}", 1)
    set_rtl(title_p)
    for run in title_p.runs:
        run.font.name = "Arial"; run.font.cs_name = "Arial"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        set_rtl_run(run)

    def meta_line(text):
        p = doc.add_paragraph(); set_rtl(p)
        run = p.add_run(text)
        run.font.name = "Arial"; run.font.cs_name = "Arial"; run.font.size = Pt(11)
        set_rtl_run(run)

    if r.get('publisher'):    meta_line(f"מפרסם: {r['publisher']}")
    if r.get('publish_date'): meta_line(f"תאריך פרסום: {r['publish_date']}")
    if r.get('deadline'):     meta_line(f"מועד הגשה: {r['deadline']}")
    if r.get('update_date'):  meta_line(f"תאריך עדכון: {r['update_date']}")
    doc.add_paragraph()

    # Parse pipe-separated questions
    rows = []
    for line in questions_raw.strip().splitlines():
        line = line.strip()
        if not line: continue
        parts = line.split("|")
        if len(parts) >= 4:
            rows.append((parts[0].strip(), parts[1].strip(), parts[2].strip(), "|".join(parts[3:]).strip()))
        elif line:
            num = str(len(rows)+1)
            rows.append((num, "כללי", "כללי", line.lstrip("0123456789. ")))

    # Column order as written in file — RTL doc renders col1 on the right
    # so מספר(col1) | עמוד(col2) | סעיף(col3) | שאלה(col4) displays correctly R→L
    headers_rtl  = ["שאלה", "סעיף", "עמוד", "מספר"]
    col_widths_rtl = [Cm(11), Cm(2.5), Cm(2), Cm(1.5)]

    def add_shd(tc, color):
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def cell_rtl_run(cell, text, bold=False, size=10, color=None):
        p = cell.paragraphs[0]; p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()
        b = OxmlElement('w:bidi'); b.set(qn('w:val'),'1'); pPr.append(b)
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'right'); pPr.append(jc)
        run = p.add_run(text)
        run.font.name = "Arial"; run.font.cs_name = "Arial"; run.font.size = Pt(size); run.bold = bold
        if color: run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'),'1'); rPr.append(rtl)
        rPr.append(OxmlElement('w:cs'))

    table = doc.add_table(rows=1+len(rows), cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, (h_text, w) in enumerate(zip(headers_rtl, col_widths_rtl)):
        cell = hdr.cells[i]; cell.width = w
        add_shd(cell._tc, '1E3A5F')
        cell_rtl_run(cell, h_text, bold=True, size=11, color=RGBColor(0xFF,0xFF,0xFF))

    # Data rows
    for ri, (num, page, section, question) in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        fill_color = 'EEF3FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, text in enumerate([question, section, page, num]):
            cell = row_cells[ci]
            add_shd(cell._tc, fill_color)
            cell_rtl_run(cell, text, size=10)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    tid = r.get("tender_id", "tender")
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''%D7%A9%D7%90%D7%9C%D7%95%D7%AA_{tid}.docx"})


# ── Email digest ──────────────────────────────────────────────────────────────

@app.post("/api/send-email")
async def send_email_digest(body: dict, u: str = Depends(auth)):
    from emailer import send_digest, build_html_digest
    from datetime import date
    import re

    tenders   = body.get("tenders", [])
    min_level = body.get("min_level", "medium")
    recipient = body.get("to", "").strip() or _load_settings()["email"]["recipient"]

    level_rank = {"high": 0, "medium": 1, "low": 2}
    threshold  = level_rank.get(min_level, 0)

    def get_rank(analysis):
        first_line = (analysis or "").split("\n")[0]
        if "גבוהה" in first_line: return 0
        if "בינונית" in first_line: return 1
        return 2

    def extract_summary(analysis: str) -> str:
        """Extract the first סיכום section from analysis."""
        lines = analysis.splitlines()
        in_section = False
        result = []
        for line in lines:
            if re.search(r'סיכום', line):
                in_section = True
                continue
            if in_section:
                if re.match(r'^#{1,3}\s', line) and result:
                    break
                if line.strip():
                    result.append(line.strip().replace('**','').replace('#',''))
        return ' '.join(result[:5]) if result else analysis[:300]

    filtered = [t for t in tenders if get_rank(t.get("analysis","")) <= threshold]
    if not filtered:
        return {"ok": False, "msg": "לא נמצאו מכרזים ברמה שנבחרה"}

    app_url = "https://tender-scanner.up.railway.app"
    def badge(a):
        if "גבוהה" in a: return "🟢 גבוהה"
        if "בינונית" in a: return "🟡 בינונית"
        return "🔴 נמוכה"
    def badge_color(a):
        if "גבוהה" in a: return "#1a7a1a"
        if "בינונית" in a: return "#b36b00"
        return "#8b0000"

    cards = []
    for t in filtered:
        analysis = t.get("analysis","")
        summary  = extract_summary(analysis)
        color    = badge_color(analysis)
        cards.append(f"""
        <div style="border:1px solid #ddd;border-radius:8px;padding:16px;margin-bottom:20px;font-family:Arial,sans-serif;direction:rtl;text-align:right;">
          <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px">
            <h2 style="margin:0;font-size:16px;color:{color}">{t.get('title','')}</h2>
            <span style="background:{color};color:#fff;padding:3px 10px;border-radius:12px;font-size:12px;white-space:nowrap">{badge(analysis)}</span>
          </div>
          <p style="margin:4px 0;color:#555;font-size:13px">
            {t.get('publisher','') or ''}
            {' | פרסום: '+t['publish_date'] if t.get('publish_date') else ''}
            {' | הגשה: '+t['deadline'] if t.get('deadline') else ''}
          </p>
          <hr style="border:none;border-top:1px solid #eee;margin:10px 0">
          <p style="font-size:14px;line-height:1.7;margin:0 0 12px">{summary}</p>
          <div style="display:flex;gap:12px">
            <a href="{t.get('url','')}" style="color:#2563EB;font-size:13px">🔗 עמוד המכרז</a>
            <a href="{app_url}" style="color:#2563EB;font-size:13px">📊 ניתוח מלא במערכת</a>
          </div>
        </div>""")

    run_date = date.today().strftime("%d/%m/%Y")
    level_label = {"high":"גבוהה בלבד","medium":"בינונית וגבוהה","low":"כל הרמות"}.get(min_level,"")
    html = f"""<html><body style="background:#f5f5f5;padding:20px">
      <h1 style="font-family:Arial,sans-serif;direction:rtl;text-align:right;color:#1a1a2e">
        סריקת מכרזים — {run_date}
      </h1>
      <p style="font-family:Arial,sans-serif;direction:rtl;text-align:right;color:#555">
        {len(filtered)} מכרזים ברמה: {level_label}
      </p>
      {''.join(cards)}
      <p style="font-family:Arial,sans-serif;font-size:12px;color:#999;text-align:center;margin-top:30px">Electra Target Tools</p>
    </body></html>"""

    resend_key = os.environ.get("RESEND_API_KEY","")
    if not resend_key:
        return {"ok": False, "msg": "RESEND_API_KEY לא מוגדר בסביבה"}

    import urllib.request, json as _json
    subject = f"[Electra Target] {len(filtered)} מכרזים — {run_date}"
    try:
        payload = _json.dumps({"from": "Electra Target <onboarding@resend.dev>", "to": [recipient], "subject": subject, "html": html}).encode()
        req = urllib.request.Request("https://api.resend.com/emails", data=payload,
            headers={"Authorization": f"Bearer {resend_key}", "Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(req, timeout=15) as resp:
            return {"ok": True, "count": len(filtered)}
    except urllib.error.HTTPError as e:
        body = e.read().decode()
        return {"ok": False, "msg": f"HTTP {e.code}: {body}"}
    except Exception as e:
        import traceback
        return {"ok": False, "msg": traceback.format_exc()[-600:]}


# ── Learning mode ─────────────────────────────────────────────────────────────

@app.post("/api/learn/analyze")
async def learn_analyze(body: dict, u: str = Depends(auth)):
    url = body.get("url", "").strip()
    settings = _load_settings()
    tid = _extract_id_from_url(url)
    meta = {"tender_id": tid, "title": url, "url": url, "publish_date": "", "update_date": ""}
    tender = await asyncio.wait_for(fetch_tender_detail(meta, settings), timeout=90)
    if not tender.title or tender.title.startswith("http"):
        tender.title = f"מכרז {tid}"
    td = {
        "tender_id": tender.tender_id, "title": tender.title, "url": tender.url,
        "publisher": tender.publisher, "deadline": tender.deadline,
        "pdf_text": tender.pdf_text, "has_pdf": bool(tender.pdf_text),
    }
    if not tender.pdf_text:
        result = {"title": tender.title, "url": url, "publisher": tender.publisher,
                  "deadline": tender.deadline, "has_pdf": False, "analysis": "NO_PDF"}
    else:
        client = _client()
        loop = asyncio.get_running_loop()
        knowledge = _load_knowledge()
        result = await loop.run_in_executor(
            None, lambda t=tender: analyze_tender(t, settings, client, knowledge=knowledge)
        )
    return {"tender": td, "result": result}

@app.post("/api/learn/reanalyze")
async def learn_reanalyze(body: dict, _: str = Depends(auth)):
    td = body.get("tender", {})
    feedback = body.get("feedback", [])
    settings = _load_settings()
    tender_obj = Tender(
        tender_id=td["tender_id"], title=td["title"], url=td["url"],
        publisher=td.get("publisher",""), deadline=td.get("deadline",""),
        pdf_text=td.get("pdf_text",""), raw_metadata={"publish_date":"","update_date":""},
    )
    client = _client()
    loop = asyncio.get_running_loop()
    knowledge = _load_knowledge()
    result = await loop.run_in_executor(
        None,
        lambda: analyze_tender(tender_obj, settings, client, knowledge=knowledge, session_feedback=feedback)
    )
    return result

@app.post("/api/learn/save")
async def learn_save(body: dict, _: str = Depends(auth)):
    title = body.get("title", "")
    feedback = body.get("feedback", [])
    client = _client()
    loop = asyncio.get_running_loop()
    new_g = await loop.run_in_executor(None, lambda: distill_knowledge(title, feedback, client))
    existing = _load_knowledge()
    existing.extend(new_g)
    _save_knowledge(existing)
    return {"new": new_g, "total": len(existing)}


# ── Shift comparison ──────────────────────────────────────────────────────────

_DEFAULT_SHIFTS_CFG = {
    "excel_columns": {"date": "B", "start_time": "C", "end_time": "D", "worker_name": "F"},
    "excel_has_header": True,
    "rules": {"gap_threshold_minutes": 30, "default_start_time": "10:00"},
    "aliases": {},
    "managers": [],
    "ignored_names": [],
}

def _shifts_cfg_path(u: str) -> Path:
    return _udir(u) / "shifts_config.json"

def _load_shifts_cfg(u: str) -> dict:
    return _rj(_shifts_cfg_path(u), dict(_DEFAULT_SHIFTS_CFG))

def _save_shifts_cfg(u: str, cfg: dict):
    _wj(_shifts_cfg_path(u), cfg)

@app.get("/api/shifts/config")
async def shifts_config(u: str = Depends(auth)):
    return _load_shifts_cfg(u)

@app.post("/api/shifts/config")
async def save_shifts_config(body: dict, u: str = Depends(auth)):
    _save_shifts_cfg(u, body)
    return {"ok": True}

@app.post("/api/shifts/validate")
async def shifts_validate(_: str = Depends(auth), message: str = Form(...)):
    return {"suspicious": find_suspicious_lines(message)}

GMAIL_EXCEL_COLUMNS = {
    "date":        "A",
    "worker_name": "D",
    "start_time":  "J",
    "end_time":    "P",
}

@app.post("/api/shifts/compare")
async def shifts_compare(
    u: str = Depends(auth),
    message: str = Form(...),
    excel: UploadFile = File(...),
    source: str = Form("upload"),
):
    cfg = _load_shifts_cfg(u)

    # Build aliases and managers from workers table (takes precedence over cfg)
    workers = _load_workers(u)
    worker_aliases = {
        w["nickname"]: w["full_name"]
        for w in workers
        if w.get("nickname") and w.get("full_name") and w["nickname"] != w["full_name"]
    }
    worker_managers = [
        w["full_name"] for w in workers if w.get("rank") == "manager" and w.get("full_name")
    ]
    known_worker_names = {w["full_name"] for w in workers if w.get("full_name")}

    compare_cfg = dict(cfg)
    merged_aliases = dict(cfg.get("aliases", {}))
    merged_aliases.update(worker_aliases)
    compare_cfg["aliases"] = merged_aliases
    if worker_managers:
        compare_cfg["managers"] = worker_managers
    compare_cfg["known_workers"] = known_worker_names

    excel_bytes = await excel.read()
    if source == "gmail":
        parse_cfg = dict(cfg)
        parse_cfg["excel_columns"] = GMAIL_EXCEL_COLUMNS
        parse_cfg["excel_has_header"] = True
    else:
        parse_cfg = cfg

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(excel_bytes)
        tmp_path = Path(tmp.name)
    try:
        msg_entries = parse_message(message)
        excel_entries = parse_excel(str(tmp_path), parse_cfg)
        results = compare(msg_entries, excel_entries, compare_cfg)

        out_path = Path(tempfile.mktemp(suffix=".xlsx"))
        export_to_excel(results, str(out_path))
        excel_bytes = out_path.read_bytes()
        out_path.unlink(missing_ok=True)

        token = secrets.token_urlsafe(16)
        _excel_cache[token] = excel_bytes

        def _f(v):
            if hasattr(v, "strftime"):
                return v.strftime("%H:%M") if hasattr(v, "hour") else v.strftime("%d/%m/%Y")
            return str(v) if v else ""

        return {
            "results": [
                {
                    "status": r["status"],
                    "worker_name": r.get("worker_name", ""),
                    "workplace": r.get("workplace", ""),
                    "date": _f(r.get("date")),
                    "start_time": _f(r.get("start_time")),
                    "end_time": _f(r.get("end_time")),
                    "hours": round(r["hours"], 2) if r.get("hours") is not None else None,
                    "sales": str(r.get("sales", "")) if r.get("sales") else "",
                    "notes": r.get("notes", ""),
                }
                for r in results
            ],
            "token": token,
        }
    finally:
        tmp_path.unlink(missing_ok=True)

@app.get("/api/shifts/download/{token}")
async def shifts_download(token: str, _: str = Depends(auth)):
    data = _excel_cache.get(token)
    if not data:
        raise HTTPException(404, "קובץ לא נמצא")
    return Response(content=data,
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=comparison.xlsx"})


# ── Saved Shifts ─────────────────────────────────────────────────────────────

def _saved_shifts_path(u: str) -> Path:
    return _udir(u) / "saved_shifts.json"

def _load_saved_shifts(u: str) -> list:
    return _rj(_saved_shifts_path(u), [])

def _save_saved_shifts(u: str, rows: list):
    _saved_shifts_path(u).write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

@app.post("/api/shifts/saved")
async def save_shifts(request: Request, u: str = Depends(auth)):
    import uuid as _uuid
    body = await request.json()
    existing = _load_saved_shifts(u)
    for row in body:
        row["id"] = str(_uuid.uuid4())
    existing.extend(body)
    _save_saved_shifts(u, existing)
    return {"ok": True, "saved": len(body)}

@app.get("/api/shifts/saved")
async def get_saved_shifts(u: str = Depends(auth), date: str = Query(None), name: str = Query(None)):
    rows = _load_saved_shifts(u)
    if date:
        rows = [r for r in rows if date in str(r.get("date", ""))]
    if name:
        nl = name.lower()
        rows = [r for r in rows if nl in str(r.get("worker_name", "")).lower()]
    rows.sort(key=lambda r: str(r.get("date", "")))
    return rows

@app.get("/api/shifts/saved/export")
async def export_saved_shifts(u: str = Depends(auth), date: str = Query(None), name: str = Query(None)):
    rows = _load_saved_shifts(u)
    if date:
        rows = [r for r in rows if date in str(r.get("date", ""))]
    if name:
        nl = name.lower()
        rows = [r for r in rows if nl in str(r.get("worker_name", "")).lower()]
    rows.sort(key=lambda r: str(r.get("date", "")))

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    STATUS_COLORS = {"ok": "C6EFCE", "gap": "FFEB9C", "missing_msg": "FFCC99", "missing_excel": "FFC7CE"}
    wb = Workbook()
    ws = wb.active
    ws.title = "שעות עובדים"
    ws.sheet_view.rightToLeft = True
    headers = ["תאריך", "שם עובד", "מקום עבודה", "שעת התחלה", "שעת סיום", "שעות", "מכירות", "הערות", "סטטוס"]
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=ci, value=h)
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", start_color="D9D9D9")
    for ri, row in enumerate(rows, 2):
        vals = [row.get("date",""), row.get("worker_name",""), row.get("workplace",""),
                row.get("start_time",""), row.get("end_time",""), row.get("hours",""),
                row.get("sales",""), row.get("notes",""), row.get("status","")]
        color = STATUS_COLORS.get(row.get("status",""), "FFFFFF")
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row=ri, column=ci, value=v)
            c.fill = PatternFill("solid", start_color=color)
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        wb.save(tmp.name)
        data = Path(tmp.name).read_bytes()
    return Response(content=data,
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=worker_hours.xlsx"})

@app.put("/api/shifts/saved/{row_id}")
async def update_saved_shift(row_id: str, body: dict, u: str = Depends(auth)):
    rows = _load_saved_shifts(u)
    for i, r in enumerate(rows):
        if r.get("id") == row_id:
            body["id"] = row_id
            rows[i] = body
            _save_saved_shifts(u, rows)
            return {"ok": True}
    raise HTTPException(404, "שורה לא נמצאה")

@app.delete("/api/shifts/saved/{row_id}")
async def delete_saved_shift(row_id: str, u: str = Depends(auth)):
    rows = _load_saved_shifts(u)
    rows = [r for r in rows if r.get("id") != row_id]
    _save_saved_shifts(u, rows)
    return {"ok": True}


# ── Workers ───────────────────────────────────────────────────────────────────

def _workers_path(u: str) -> Path:
    return _udir(u) / "workers.json"

def _load_workers(u: str) -> list:
    return _rj(_workers_path(u), [])

def _save_workers(u: str, workers: list):
    _workers_path(u).write_text(json.dumps(workers, ensure_ascii=False, indent=2), encoding="utf-8")

@app.post("/api/workers/upload")
async def upload_workers(u: str = Depends(auth), file: UploadFile = File(...)):
    import uuid as _uuid
    import pandas as pd
    data = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        df = pd.read_excel(tmp_path, header=0)
        df.columns = [str(c).strip() for c in df.columns]

        # Map column names to internal fields — partial case-insensitive match
        keywords = {
            "full_name":    ["שם", "name"],
            "id_number":    ["ז", "id", "עובד", "מספר"],
            "nickname":     ["כינוי", "nick"],
            "manager":      ["מנהל", "manager"],
            "rank":         ["דרגה", "תפקיד", "type", "rank"],
            "sales_target": ["יעד", "מכירות", "target", "sales"],
        }
        # Build field→column mapping
        field_col: dict[str, str] = {}
        for field, kws in keywords.items():
            for col in df.columns:
                col_l = col.lower()
                if any(kw.lower() in col_l for kw in kws) and field not in field_col:
                    field_col[field] = col
                    break

        # Positional fallback: name, id, nickname, manager, rank, sales_target
        positional = ["full_name", "id_number", "nickname", "manager", "rank", "sales_target"]
        for i, field in enumerate(positional):
            if field not in field_col and i < len(df.columns):
                field_col[field] = df.columns[i]

        def cell(row, field):
            col = field_col.get(field)
            if col is None:
                return ""
            val = row[col]
            return "" if (val is None or (isinstance(val, float) and pd.isna(val))) else str(val).strip()

        workers = []
        for _, row in df.iterrows():
            full_name = cell(row, "full_name")
            if not full_name:
                continue
            rank_val = cell(row, "rank").lower()
            rank = "manager" if any(x in rank_val for x in ["מנהל", "manager"]) else "worker"
            workers.append({
                "id":           str(_uuid.uuid4()),
                "full_name":    full_name,
                "id_number":    cell(row, "id_number"),
                "nickname":     cell(row, "nickname"),
                "manager":      cell(row, "manager"),
                "rank":         rank,
                "sales_target": cell(row, "sales_target"),
            })
        _save_workers(u, workers)
        return {"ok": True, "count": len(workers)}
    finally:
        Path(tmp_path).unlink(missing_ok=True)

@app.get("/api/workers")
async def get_workers(u: str = Depends(auth)):
    return _load_workers(u)

@app.post("/api/workers")
async def add_worker(body: dict, u: str = Depends(auth)):
    import uuid as _uuid
    workers = _load_workers(u)
    body["id"] = str(_uuid.uuid4())
    workers.append(body)
    _save_workers(u, workers)
    return {"ok": True, "worker": body}

@app.put("/api/workers/{worker_id}")
async def update_worker(worker_id: str, body: dict, u: str = Depends(auth)):
    workers = _load_workers(u)
    for i, w in enumerate(workers):
        if w.get("id") == worker_id:
            body["id"] = worker_id
            workers[i] = body
            _save_workers(u, workers)
            return {"ok": True}
    raise HTTPException(404, "עובד לא נמצא")

@app.delete("/api/workers/{worker_id}")
async def delete_worker(worker_id: str, u: str = Depends(auth)):
    workers = _load_workers(u)
    workers = [w for w in workers if w.get("id") != worker_id]
    _save_workers(u, workers)
    return {"ok": True}


# ── Sales ─────────────────────────────────────────────────────────────────────

def _sales_path(u: str) -> Path:
    return _udir(u) / "sales.json"

def _load_sales(u: str) -> list:
    return _rj(_sales_path(u), [])

def _save_sales(u: str, sales: list):
    _sales_path(u).write_text(json.dumps(sales, ensure_ascii=False, indent=2), encoding="utf-8")

@app.get("/api/sales")
async def get_sales(u: str = Depends(auth)):
    return _load_sales(u)

@app.post("/api/sales/upload")
async def upload_sales(u: str = Depends(auth), file: UploadFile = File(...)):
    import pandas as pd
    contents = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(contents)
        tmp_path = Path(tmp.name)
    try:
        target_sheet = "פירוט בקשות מעודכן"
        xl = pd.ExcelFile(str(tmp_path))
        sheet_name = None
        for s in xl.sheet_names:
            if s.strip() == target_sheet:
                sheet_name = s
                break
        if sheet_name is None:
            raise HTTPException(status_code=400, detail=f"גיליון '{target_sheet}' לא נמצא בקובץ")

        df = pd.read_excel(str(tmp_path), sheet_name=sheet_name, header=0, dtype=str)
        df = df.fillna("")

        def cell(row, col_idx):
            if col_idx < len(row):
                v = str(row.iloc[col_idx]).strip()
                return "" if v in ("nan", "None") else v
            return ""

        sales = []
        for _, row in df.iterrows():
            sale_num = cell(row, 0)   # A
            if not sale_num:
                continue
            date_val = cell(row, 1)   # B
            branch   = cell(row, 2)   # C
            first_name = cell(row, 3) # D
            last_name  = cell(row, 4) # E
            standing_order_raw = cell(row, 10)  # K
            standing_order = standing_order_raw == "מולא"
            revolving_l    = cell(row, 11) == "1"  # L
            revolving_m    = cell(row, 12) == "1"  # M
            revolving_h    = cell(row, 13) == "1"  # N
            revolving_xl   = cell(row, 14) == "1"  # O
            status_raw = cell(row, 15)              # P
            approved = status_raw != "תעודה מזהה לא בתוקף" and status_raw != ""

            sales.append({
                "sale_number":    sale_num,
                "date":           date_val,
                "branch":         branch,
                "first_name":     first_name,
                "last_name":      last_name,
                "standing_order": standing_order,
                "revolving_1500": revolving_l,
                "revolving_2500": revolving_m,
                "revolving_4000": revolving_h,
                "revolving_4001": revolving_xl,
                "approved":       approved,
                "status_raw":     status_raw,
            })
        _save_sales(u, sales)
        return {"ok": True, "count": len(sales)}
    finally:
        tmp_path.unlink(missing_ok=True)


@app.get("/api/report/export")
async def export_report(u: str = Depends(auth)):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    workers    = _load_workers(u)
    shifts_all = _load_saved_shifts(u)
    sales_all  = _load_sales(u)

    # Shift stats
    shift_stats: dict = {}
    for s in shifts_all:
        name = (s.get("worker_name") or "").strip()
        if not name:
            continue
        if name not in shift_stats:
            shift_stats[name] = {"dates": set(), "hours": 0.0}
        if s.get("date"):
            shift_stats[name]["dates"].add(s["date"])
        shift_stats[name]["hours"] += float(s.get("hours") or 0)

    # Sales stats
    def init_sale():
        return {"total": 0, "rev1500": 0, "rev2500": 0, "rev4000": 0, "so": 0, "issued": 0}
    sale_stats: dict = {}
    for s in sales_all:
        name = ((s.get("first_name") or "") + " " + (s.get("last_name") or "")).strip()
        if not name:
            continue
        if name not in sale_stats:
            sale_stats[name] = init_sale()
        st = sale_stats[name]
        if s.get("approved"):      st["total"]  += 1
        if s.get("revolving_1500"): st["rev1500"] += 1
        if s.get("revolving_2500"): st["rev2500"] += 1
        if s.get("revolving_4000"): st["rev4000"] += 1
        if s.get("standing_order"): st["so"]      += 1
        if (s.get("status_raw") or "").strip() == "הונפק": st["issued"] += 1

    # Group by manager
    manager_map: dict = {}
    no_manager: list = []
    for w in workers:
        if w.get("rank") == "manager":
            nm = w.get("full_name", "")
            if nm not in manager_map:
                manager_map[nm] = {"mgr": w, "workers": []}
            else:
                manager_map[nm]["mgr"] = w
    for w in workers:
        if w.get("rank") == "manager":
            continue
        mgr = (w.get("manager") or "").strip()
        if mgr and mgr in manager_map:
            manager_map[mgr]["workers"].append(w)
        elif mgr:
            if mgr not in manager_map:
                manager_map[mgr] = {"mgr": None, "workers": []}
            manager_map[mgr]["workers"].append(w)
        else:
            no_manager.append(w)

    HEADERS = [
        "ת.ז", "שם עובד", "מנהל", "משמרות", "שעות",
        "סה\"כ מכירות", "אשראי עד 1500", "אשראי 1501-2500", "אשראי 2501-4000",
        "סה\"כ הו\"ק", "סה\"כ ארנוקים", "ממוצע מכירות/שעה",
        "צפי שעות", "יעד", "צפי מכירות", "% הגעה ליעד", "הנפקות",
    ]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "דוח מלא"
    ws.sheet_view.rightToLeft = True

    hdr_fill   = PatternFill("solid", fgColor="2F5496")
    mgr_fill   = PatternFill("solid", fgColor="D6DCE4")
    sum_fill   = PatternFill("solid", fgColor="BDD7EE")
    hdr_font   = Font(bold=True, color="FFFFFF", size=10)
    bold_font  = Font(bold=True, size=10)
    norm_font  = Font(size=10)
    center     = Alignment(horizontal="center")

    ws.append(HEADERS)
    for cell in ws[1]:
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = center

    def fmt(n):
        if n is None: return ""
        return round(n, 2) if n % 1 else int(n)

    def worker_data(w, is_mgr):
        name = w.get("full_name", "")
        ss   = shift_stats.get(name, {"dates": set(), "hours": 0.0})
        sa   = sale_stats.get(name, init_sale())
        shifts_n = len(ss["dates"])
        hours_n  = ss["hours"]
        avg = fmt(sa["total"] / hours_n) if hours_n else ""
        return [
            w.get("id_number", ""),
            name,
            "" if is_mgr else (w.get("manager") or ""),
            shifts_n,
            fmt(hours_n),
            sa["total"], sa["rev1500"], sa["rev2500"], sa["rev4000"],
            sa["so"], "",
            avg,
            "", w.get("sales_target", ""), "", "", sa["issued"],
        ]

    def sum_data(mgr_name, name_list):
        shifts_n = hours_n = total = rev1500 = rev2500 = rev4000 = so = issued = 0
        for name in name_list:
            ss = shift_stats.get(name, {"dates": set(), "hours": 0.0})
            sa = sale_stats.get(name, init_sale())
            shifts_n += len(ss["dates"])
            hours_n  += ss["hours"]
            total    += sa["total"]
            rev1500  += sa["rev1500"]
            rev2500  += sa["rev2500"]
            rev4000  += sa["rev4000"]
            so       += sa["so"]
            issued   += sa["issued"]
        avg = fmt(total / hours_n) if hours_n else ""
        return [
            f"סה\"כ {mgr_name}", f"סה\"כ {mgr_name}", "",
            shifts_n, fmt(hours_n),
            total, rev1500, rev2500, rev4000,
            so, "", avg,
            "", "", "", "", issued,
        ]

    def append_worker(row_data, is_mgr=False, is_sum=False):
        ws.append(row_data)
        row = ws.max_row
        for cell in ws[row]:
            cell.font = bold_font if (is_mgr or is_sum) else norm_font
            if is_sum:
                cell.fill = sum_fill
            elif is_mgr:
                cell.fill = mgr_fill
            if cell.column > 3:
                cell.alignment = center

    for mgr_name, group in manager_map.items():
        for w in group["workers"]:
            append_worker(worker_data(w, False))
        if group["mgr"]:
            append_worker(worker_data(group["mgr"], True), is_mgr=True)
        all_names = [w.get("full_name","") for w in group["workers"]]
        if group["mgr"]:
            all_names.append(group["mgr"].get("full_name",""))
        append_worker(sum_data(mgr_name, all_names), is_sum=True)

    for w in no_manager:
        append_worker(worker_data(w, False))

    for i, col in enumerate(ws.columns, 1):
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 28)

    out = Path(tempfile.mktemp(suffix=".xlsx"))
    wb.save(str(out))
    data = out.read_bytes()
    out.unlink(missing_ok=True)

    from fastapi.responses import Response
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="worker_report.xlsx"'},
    )


# ── Recruiter Analysis ────────────────────────────────────────────────────────

RECRUITER_CONFIG_FILE = BASE_DIR / "data" / "shared" / "recruiter_config.json"
RECRUITER_DATA_FILE = BASE_DIR / "data" / "shared" / "recruiter_data.json"

_DEFAULT_RECRUITER_CONFIG: dict = {
    "recruiters": [],
    "long_call_threshold_minutes": 8,
    "repeat_call_threshold": 2,
    "default_days_back": 7,
}

def _load_recruiter_config() -> dict:
    return _rj(RECRUITER_CONFIG_FILE, dict(_DEFAULT_RECRUITER_CONFIG))

def _save_recruiter_config(cfg: dict):
    _wj(RECRUITER_CONFIG_FILE, cfg)

# ── WhatsApp Webhook ──────────────────────────────────────────────────────────
WHATSAPP_VERIFY_TOKEN = os.environ.get("WHATSAPP_VERIFY_TOKEN", "electra_target_verify_2024")
_whatsapp_messages: list = []  # stores last received shift messages

@app.get("/api/whatsapp/webhook")
async def whatsapp_verify(request: Request):
    """Meta webhook verification handshake."""
    params = dict(request.query_params)
    if params.get("hub.verify_token") == WHATSAPP_VERIFY_TOKEN and params.get("hub.mode") == "subscribe":
        return Response(content=params.get("hub.challenge", ""), media_type="text/plain")
    raise HTTPException(403, "Invalid verify token")

@app.post("/api/whatsapp/webhook")
async def whatsapp_receive(request: Request):
    """Receive incoming WhatsApp messages."""
    data = await request.json()
    try:
        for entry in data.get("entry", []):
            for change in entry.get("changes", []):
                messages = change.get("value", {}).get("messages", [])
                for msg in messages:
                    if msg.get("type") == "text":
                        text = msg["text"]["body"]
                        if text.strip().startswith("!משמרות"):
                            content = text[len("!משמרות"):].strip()
                            _whatsapp_messages.insert(0, {"text": content, "timestamp": msg.get("timestamp","")})
                            _whatsapp_messages[:] = _whatsapp_messages[:10]  # keep last 10
    except Exception:
        pass
    return {"status": "ok"}

@app.get("/api/whatsapp/latest-message")
async def whatsapp_latest(u: str = Depends(auth)):
    if not _whatsapp_messages:
        return {"ok": False, "msg": "לא נמצאה הודעת משמרות"}
    return {"ok": True, "text": _whatsapp_messages[0]["text"]}


# ── Gmail OAuth ───────────────────────────────────────────────────────────────
GMAIL_CLIENT_ID     = os.environ.get("GOOGLE_CLIENT_ID", "")
GMAIL_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
GMAIL_REDIRECT_URI  = "https://tender-scanner.up.railway.app/auth/gmail/callback"
GMAIL_SCOPES        = "https://www.googleapis.com/auth/gmail.readonly"
_GMAIL_TOKENS_FILE = BASE_DIR / "data" / "shared" / "gmail_tokens.json"

def _load_gmail_tokens() -> dict:
    return _rj(_GMAIL_TOKENS_FILE, {})

def _save_gmail_token(u: str, tokens: dict):
    all_tokens = _load_gmail_tokens()
    all_tokens[u] = tokens
    _GMAIL_TOKENS_FILE.write_text(json.dumps(all_tokens, ensure_ascii=False, indent=2), encoding="utf-8")

@app.get("/auth/gmail/connect")
async def gmail_connect(u: str = Depends(auth)):
    import urllib.parse
    params = urllib.parse.urlencode({
        "client_id": GMAIL_CLIENT_ID,
        "redirect_uri": GMAIL_REDIRECT_URI,
        "response_type": "code",
        "scope": GMAIL_SCOPES,
        "access_type": "offline",
        "prompt": "consent",
        "state": u,
    })
    from fastapi.responses import RedirectResponse
    return RedirectResponse(f"https://accounts.google.com/o/oauth2/v2/auth?{params}")

@app.get("/auth/gmail/callback")
async def gmail_callback(code: str, state: str):
    import urllib.request, urllib.parse, json as _json
    data = urllib.parse.urlencode({
        "code": code,
        "client_id": GMAIL_CLIENT_ID,
        "client_secret": GMAIL_CLIENT_SECRET,
        "redirect_uri": GMAIL_REDIRECT_URI,
        "grant_type": "authorization_code",
    }).encode()
    req = urllib.request.Request("https://oauth2.googleapis.com/token", data=data,
        headers={"Content-Type": "application/x-www-form-urlencoded"}, method="POST")
    with urllib.request.urlopen(req) as resp:
        tokens = _json.loads(resp.read())
    _save_gmail_token(state, tokens)
    from fastapi.responses import HTMLResponse
    return HTMLResponse("<script>window.close();window.opener&&window.opener.postMessage('gmail_connected','*')</script>✅ Gmail מחובר! ניתן לסגור חלון זה.")

@app.get("/auth/gmail/status")
async def gmail_status(u: str = Depends(auth)):
    return {"connected": u in _load_gmail_tokens()}

@app.post("/api/shifts/fetch-from-gmail")
async def fetch_shifts_from_gmail(body: dict = {}, u: str = Depends(auth)):
    import urllib.request, urllib.parse, json as _json, base64
    tokens = _load_gmail_tokens().get(u)
    if not tokens:
        raise HTTPException(400, "Gmail לא מחובר")
    access_token = tokens.get("access_token", "")

    def gmail_get(url):
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {access_token}"})
        with urllib.request.urlopen(req) as r:
            return _json.loads(r.read())

    # Search for email from clock2go, optionally filtered by date in subject
    date_str = body.get("date", "")
    subject_query = f'דו"ח נוכחות כולל משימות יומי {date_str}'.strip()
    q = urllib.parse.quote(f'from:support@clock2go.co.il subject:{subject_query}')
    result = gmail_get(f"https://gmail.googleapis.com/gmail/v1/users/me/messages?q={q}&maxResults=1")
    messages = result.get("messages", [])
    if not messages:
        raise HTTPException(404, "לא נמצא מייל מ-clock2go עם קובץ נוכחות")

    msg = gmail_get(f"https://gmail.googleapis.com/gmail/v1/users/me/messages/{messages[0]['id']}")
    subject = next((h["value"] for h in msg["payload"]["headers"] if h["name"] == "Subject"), "")

    # Find Excel attachment
    def find_parts(part):
        if part.get("filename","").endswith((".xlsx",".xls")) and part.get("body",{}).get("attachmentId"):
            return part
        for p in part.get("parts", []):
            found = find_parts(p)
            if found:
                return found
        return None

    att_part = find_parts(msg["payload"])
    if not att_part:
        raise HTTPException(404, "לא נמצא קובץ Excel במייל")

    att_id = att_part["body"]["attachmentId"]
    att = gmail_get(f"https://gmail.googleapis.com/gmail/v1/users/me/messages/{messages[0]['id']}/attachments/{att_id}")
    file_bytes = base64.urlsafe_b64decode(att["data"])

    # Save to temp and return as upload token
    token = secrets.token_hex(16)
    _excel_cache[token] = file_bytes
    return {"ok": True, "token": token, "filename": att_part["filename"], "subject": subject}


@app.get("/api/recruiter/config")
async def recruiter_config_get(_: str = Depends(auth)):
    return _load_recruiter_config()

@app.post("/api/recruiter/config")
async def recruiter_config_post(body: dict, _: str = Depends(auth)):
    _save_recruiter_config(body)
    return {"ok": True}

@app.post("/api/recruiter/upload")
async def recruiter_upload(excel: UploadFile = File(...), _: str = Depends(auth)):
    import pandas as pd
    from datetime import datetime as _dt

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(await excel.read())
        tmp_path = Path(tmp.name)

    try:
        fname = excel.filename or ""
        if fname.lower().endswith(".csv"):
            df = pd.read_csv(str(tmp_path))
        else:
            df = pd.read_excel(str(tmp_path), sheet_name="פיד", header=0)
    finally:
        tmp_path.unlink(missing_ok=True)

    def _ext(val) -> str | None:
        try:
            s = str(int(float(val)))
            if s.startswith("910") and len(s) >= 6:
                return s[3:].lstrip("0") or s[3:]
            if 2 <= len(s) <= 4:
                return s.lstrip("0") or s
        except Exception:
            pass
        return None

    calls = []
    for _, row in df.iterrows():
        ext = _ext(row.get("src"))
        if not ext:
            continue
        raw_date = row.get("calldate")
        try:
            if pd.isna(raw_date):
                continue
        except Exception:
            if not raw_date:
                continue
        try:
            dt = pd.to_datetime(raw_date)
            date_str = dt.strftime("%Y-%m-%d")
            time_str = dt.strftime("%H:%M")
            hour = int(dt.hour)
        except Exception:
            continue
        try:
            duration = int(float(row.get("billsec") or 0))
        except Exception:
            duration = 0
        answered = str(row.get("disposition", "")).upper() == "ANSWERED"
        dst = str(row.get("dst", "")).rstrip(".0").strip()
        calls.append({
            "date": date_str, "time": time_str, "hour": hour,
            "extension": ext, "dst": dst,
            "duration_sec": duration, "answered": answered,
        })

    _wj(RECRUITER_DATA_FILE, {"last_updated": _dt.now().isoformat(), "calls": calls})
    return {"ok": True, "total_calls": len(calls)}


@app.get("/api/recruiter/data")
async def recruiter_data(
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    _: str = Depends(auth),
):
    import collections

    raw = _rj(RECRUITER_DATA_FILE, None)
    if not raw:
        return {"last_updated": None, "recruiters": [], "all_dates": [], "full_range": None}

    cfg = _load_recruiter_config()
    recruiter_map = {r["extension"]: r["name"] for r in cfg.get("recruiters", [])}
    long_sec = cfg.get("long_call_threshold_minutes", 8) * 60
    repeat_min = cfg.get("repeat_call_threshold", 2)

    all_calls = raw.get("calls", [])
    all_call_dates = sorted(set(c["date"] for c in all_calls))
    full_range = {"from": all_call_dates[0], "to": all_call_dates[-1]} if all_call_dates else None

    df = date_from or (all_call_dates[0] if all_call_dates else "")
    dt = date_to or (all_call_dates[-1] if all_call_dates else "")
    calls = [c for c in all_calls if df <= c["date"] <= dt]
    all_dates = sorted(set(c["date"] for c in calls))

    by_ext: dict = collections.defaultdict(list)
    for c in calls:
        by_ext[c["extension"]].append(c)

    recruiters = []
    for ext, rcalls in by_ext.items():
        name = recruiter_map.get(ext, f"שלוחה {ext}")
        answered = [c for c in rcalls if c["answered"]]
        total = len(rcalls)
        ans_count = len(answered)
        ans_rate = round(ans_count / total * 100, 1) if total else 0
        total_sec = sum(c["duration_sec"] for c in answered)
        total_min = round(total_sec / 60, 1)
        avg_dur = round(total_sec / ans_count / 60, 1) if ans_count else 0
        work_days = len(set(c["date"] for c in rcalls))
        avg_per_day = round(total / work_days, 1) if work_days else 0
        times = sorted(c["time"] for c in rcalls)

        hourly: dict = collections.defaultdict(float)
        for c in answered:
            hourly[c["hour"]] += c["duration_sec"] / 60
        hourly_dist = {str(h): round(hourly.get(h, 0), 1) for h in range(8, 18)}

        daily_calls: dict = collections.defaultdict(int)
        daily_minutes: dict = collections.defaultdict(float)
        for c in rcalls:
            daily_calls[c["date"]] += 1
        for c in answered:
            daily_minutes[c["date"]] += c["duration_sec"] / 60

        long_calls = sorted(
            [{"date": c["date"], "time": c["time"], "dst": c["dst"],
              "minutes": round(c["duration_sec"] / 60, 1)}
             for c in answered if c["duration_sec"] >= long_sec],
            key=lambda x: x["minutes"], reverse=True,
        )

        dst_counts = collections.Counter(c["dst"] for c in rcalls if c["dst"])
        repeat_numbers = [
            {"number": num, "count": cnt}
            for num, cnt in dst_counts.most_common(20)
            if cnt >= repeat_min
        ]

        recruiters.append({
            "name": name, "extension": ext,
            "total_calls": total, "answered_calls": ans_count, "answer_rate": ans_rate,
            "total_minutes": total_min, "avg_duration_minutes": avg_dur,
            "work_days": work_days, "avg_calls_per_day": avg_per_day,
            "first_call": times[0] if times else None,
            "last_call": times[-1] if times else None,
            "hourly_distribution": hourly_dist,
            "daily_calls": dict(daily_calls),
            "daily_minutes": {d: round(v, 1) for d, v in daily_minutes.items()},
            "long_calls": long_calls,
            "repeat_numbers": repeat_numbers,
        })

    known_order = {r["extension"]: i for i, r in enumerate(cfg.get("recruiters", []))}
    recruiters.sort(key=lambda r: known_order.get(r["extension"], 999))

    return {"last_updated": raw.get("last_updated"), "all_dates": all_dates, "full_range": full_range, "recruiters": recruiters}
